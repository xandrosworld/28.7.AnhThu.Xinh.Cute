"""Create the completed Chapter 3 report without modifying the source DOCX.

Usage after verification:
    python docs/generate_completion_report.py \
        --test-result "78 passed, 1 SQL Server concurrency test skipped locally" \
        --coverage "86.58%" \
        --benchmark-max-ms "71.778" \
        --sqlserver pending \
        --browser chromium-local

The result is always written under ``docs/``.  Test results are required so a
stale or invented count cannot silently enter the submission report.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = (
    ROOT
    / "docs"
    / "49K21.1_NguyenHoangThanhTruc_Chuong3_KPI4_Goc.docx"
)
OUTPUT = (
    ROOT
    / "docs"
    / "49K21.1_NguyenHoangThanhTruc_Chuong3_KPI4_HoanThien.docx"
)
def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--test-result",
        required=True,
        help='Kết quả pytest thật, ví dụ "78 passed"',
    )
    parser.add_argument(
        "--coverage",
        default="Chưa ghi nhận",
        help='Coverage thật, ví dụ "86.58%%"; không cung cấp nếu chưa đo',
    )
    parser.add_argument(
        "--benchmark-max-ms",
        required=True,
        help='Max thật từ performance_results.json, ví dụ "71.778"',
    )
    parser.add_argument(
        "--sqlserver",
        choices=("passed", "pending"),
        default="pending",
        help="Chỉ chọn passed sau khi job SQL Server thật đã xanh",
    )
    parser.add_argument(
        "--sqlserver-result",
        default="",
        help='Kết quả SQL Server thật, ví dụ "77 passed, 2 skipped"',
    )
    parser.add_argument(
        "--browser",
        choices=("pending", "chromium-local", "passed"),
        default="chromium-local",
        help="Chỉ chọn passed sau khi Chromium và Firefox CI đều xanh",
    )
    parser.add_argument("--ci-run-url", default="", help="URL run CI xanh của revision hiện tại")
    parser.add_argument("--ci-sha", default="", help="Full SHA tương ứng với run CI")
    return parser.parse_args()


def main():
    args = parse_args()
    if (
        args.sqlserver == "passed" or args.browser == "passed"
    ) and not (args.ci_run_url and args.ci_sha):
        raise ValueError(
            "--ci-run-url và --ci-sha là bắt buộc khi đánh dấu SQL Server "
            "hoặc browser đã đạt"
        )
    if not SOURCE.exists():
        raise FileNotFoundError(f"Không tìm thấy tài liệu nguồn: {SOURCE}")

    document = Document(SOURCE)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("PHỤ LỤC ĐỐI CHIẾU SẢN PHẨM WMS HOÀN THIỆN", level=1)
    document.add_paragraph(
        "Phụ lục được tạo từ bản tích hợp của ba nhánh Anh_Thu, Le_Thao và "
        "Thanh_Truc. Nội dung gốc được giữ nguyên và file nguồn không bị ghi đè."
    )
    document.add_paragraph(
        "Repository bàn giao chính thức: "
        "https://github.com/xandrosworld/28.7.AnhThu.Xinh.Cute"
    )

    document.add_heading("1. Phân công và lịch sử đóng góp", level=2)
    add_table(
        document,
        ("Thành viên", "Vai trò", "Bằng chứng"),
        (
            (
                "Nguyễn Hoàng Thanh Trúc",
                "BA",
                "Yêu cầu, luồng nghiệp vụ, acceptance, báo cáo — nhánh Thanh_Truc",
            ),
            (
                "Dương Thị Anh Thư",
                "FE Developer",
                "Design system, responsive, accessibility — nhánh Anh_Thu",
            ),
            (
                "Lê Phương Thảo",
                "BE Developer",
                "Database, API, transaction, tồn kho — nhánh Le_Thao",
            ),
        ),
    )

    document.add_heading("2. Kiến trúc và phạm vi hoàn thiện", level=2)
    for item in (
        "Một Flask modular monolith; API, service, model/repository và giao diện dùng chung.",
        "SQLAlchemy/Alembic; SQLite chạy ngay và SQL Server cấu hình qua DATABASE_URL.",
        "RBAC ADMIN/CS/WAREHOUSE; ADMIN cấu hình vai trò và đơn vị tính tại /settings.",
        "Nhập kho bắt buộc inspection trước confirm; accepted/rejected có lý do.",
        "Xuất kho kiểm email hợp đồng, start-picking/reject và picking FEFO/FIFO.",
        "Decimal end-to-end, lot/pallet, stock movement, snapshot, transaction và idempotency.",
        "Dashboard 6 KPI, báo cáo, CSV, bản in và backup/restore SQLite.",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("3. Kết quả kiểm thử đã xác nhận", level=2)
    sql_result = (
        "Đạt trên SQL Server 2022 + ODBC 18: "
        f"{args.sqlserver_result or 'đã chạy đầy đủ'}; {args.ci_run_url}, SHA {args.ci_sha}"
        if args.sqlserver == "passed"
        else "Revision hiện tại chờ kết quả SQL Server 2022; chưa tuyên bố đạt"
    )
    browser_result = {
        "pending": "18 Playwright cases đã khai báo; chưa có browser result được xác minh",
        "chromium-local": (
            "18 Playwright project-cases; Chromium local 9/9 tại "
            "1366/1024/390 px; Firefox và CI pending"
        ),
        "passed": (
            f"Chromium và Firefox CI đạt tại 1366/1024/390 px; "
            f"{args.ci_run_url}, SHA {args.ci_sha}"
        ),
    }[args.browser]
    ci_result = (
        f"{args.ci_run_url} tại SHA {args.ci_sha}"
        if args.ci_run_url and args.ci_sha
        else "Revision hiện tại chưa có run CI xanh được ghi nhận"
    )
    sql_traceability = (
        "AT-18 và NFR lưu trữ/toàn vẹn trên SQL Server đã đạt tại "
        f"{args.ci_run_url}, SHA {args.ci_sha}: "
        f"{args.sqlserver_result or 'test SQL Server đã đạt'}. "
        if args.sqlserver == "passed"
        else "AT-18 chưa được đánh dấu đạt do chưa có bằng chứng CI SQL Server. "
    )
    if args.sqlserver == "passed":
        document.add_paragraph(
            "Bằng chứng CI chính thức: "
            f"{args.ci_run_url} tại SHA {args.ci_sha}. "
            "Chỉ ghi các job đạt theo kết quả của chính run này."
        )
    benchmark_display = args.benchmark_max_ms.replace(".", ",")
    add_table(
        document,
        ("Hạng mục", "Kết quả"),
        (
            ("Pytest trên máy xác minh", args.test_result),
            ("Coverage package app", args.coverage),
            (
                "NFR-005 hiệu năng",
                f"Đạt; request chậm nhất {benchmark_display} ms khi xuất CSV 5.000 movement",
            ),
            ("SQLite", "Migrate, seed và test trong quy trình xác minh"),
            ("SQL Server", sql_result),
            (
                "GitHub Actions",
                ci_result,
            ),
            (
                "Playwright",
                browser_result,
            ),
        ),
    )

    document.add_heading("4. Truy vết yêu cầu", level=2)
    document.add_paragraph(
        "BR-001 đến BR-015 và NFR-001 đến NFR-010 được truy vết tại "
        "docs/REQUIREMENTS_TRACEABILITY.md tới API/màn hình và acceptance test. "
        "NFR-005 đã đạt benchmark độc lập với 1.012 sản phẩm, 5.010 lot và "
        f"5.000 stock movement; request chậm nhất {benchmark_display} ms, thấp hơn ngưỡng "
        "5 giây. "
        f"{sql_traceability}"
        f"Browser: {browser_result}. Camera/USB/in vật lý vẫn theo checklist thủ công."
    )

    document.add_heading("5. Nghiệm thu nhanh", level=2)
    for index, step in enumerate(
        (
            "Cài dependency và tạo file cấu hình từ .env.example.",
            "Chạy migration, seed dữ liệu demo và khởi động ứng dụng theo README.",
            "Demo lần lượt tài khoản CS, Warehouse và Admin.",
            "Chạy toàn bộ pytest và xem workflow CI.",
            "Trình bày ma trận truy vết và lịch sử ba nhánh.",
        ),
        1,
    ):
        # The supplied Vietnamese Word template has no English "List Number"
        # style. Explicit numbering preserves compatibility with that template.
        document.add_paragraph(f"{index}. {step}")

    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"].font.size = Pt(13)
    document.core_properties.title = (
        "Chương 3 KPI4 — Bản hoàn thiện hệ thống WMS tích hợp"
    )
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
